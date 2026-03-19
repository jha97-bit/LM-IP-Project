import bootstrap  # noqa: F401

import io
from typing import Dict, List, Tuple

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib import colors as mcolors
import numpy as np
import pandas as pd
import streamlit as st
from sqlalchemy import text

from app.app_context import guard_page, sync_method_from_scenario
from app.sidebar_nav import render_sidebar
from core.topsis import compute_topsis
from persistence.engine import get_engine
from persistence.repositories.measurement_repo import MeasurementRepo
from persistence.repositories.preference_repo import PreferenceRepo
from persistence.repositories.result_repo import ResultRepo
from persistence.repositories.topsis_read_repo import TopsisReadRepo
from services.vft_service import VFTService

st.set_page_config(page_title="MCDA — Report Builder", layout="wide")
st.title("Step 6: Report Builder")

guard_page("pages/6_report_builder.py", require_scenario=True)
sync_method_from_scenario()
render_sidebar("pages/6_report_builder.py")

engine = get_engine()
meas_repo = MeasurementRepo(engine)
pref_repo = PreferenceRepo(engine)
result_repo = ResultRepo(engine)
topsis_read = TopsisReadRepo(engine)
vft_svc = VFTService(engine)

scenario_id = st.session_state.get("scenario_id")
user_name = st.session_state.get("user_name", "")
method_choice = st.session_state.get("method_choice", "topsis")

if not scenario_id:
    st.warning("No scenario selected. Go to Step 1.")
    if st.button("← Go to Step 1", key="rpt_goto_step1"):
        st.switch_page("pages/1_decision_setup.py")
    st.stop()

nav_left, nav_right = st.columns(2)
with nav_left:
    if st.button("← Step 5: Sensitivity", key="rpt_nav_back"):
        st.switch_page("pages/5_sensitivity.py")
with nav_right:
    if st.button("Next: History & Logs →", type="primary", key="rpt_nav_next"):
        st.switch_page("pages/7_history.py")

st.caption("Build a DOCX report with section-wise tables, charts, legends, sandbox views, comparison views, and notes.")
st.divider()


def load_decisions():
    with engine.begin() as conn:
        return [dict(r) for r in conn.execute(
            text("SELECT decision_id::text, title FROM decisions ORDER BY created_at DESC LIMIT 200")
        ).mappings().all()]


def load_scenarios(did):
    with engine.begin() as conn:
        return [dict(r) for r in conn.execute(
            text("SELECT scenario_id::text, name, method_type FROM scenarios WHERE decision_id=:did ORDER BY created_at DESC"),
            {"did": did},
        ).mappings().all()]


def load_prefs(sid):
    with engine.begin() as conn:
        return [dict(r) for r in conn.execute(
            text("SELECT preference_set_id::text, name FROM preference_sets WHERE scenario_id=:sid ORDER BY created_at DESC"),
            {"sid": sid},
        ).mappings().all()]


def load_runs(sid, pid):
    with engine.begin() as conn:
        return [dict(r) for r in conn.execute(
            text(
                """
                SELECT run_id::text, executed_at, executed_by, method, run_label
                FROM runs
                WHERE scenario_id=:sid AND preference_set_id=:pid
                ORDER BY executed_at DESC LIMIT 200
                """
            ),
            {"sid": sid, "pid": pid},
        ).mappings().all()]


def load_latest_run_for_pref(sid: str, pid: str, method: str) -> str | None:
    with engine.begin() as conn:
        row = conn.execute(
            text(
                """
                SELECT run_id::text AS run_id
                FROM runs
                WHERE scenario_id=:sid AND preference_set_id=:pid AND method=:method
                ORDER BY executed_at DESC
                LIMIT 1
                """
            ),
            {"sid": sid, "pid": pid, "method": method},
        ).mappings().first()
    return row["run_id"] if row else None


def load_criteria_meta(sid: str) -> List[dict]:
    with engine.begin() as conn:
        rows = conn.execute(
            text(
                """
                SELECT criterion_id::text AS criterion_id, name, direction, scale_type
                FROM criteria
                WHERE scenario_id = :sid
                ORDER BY created_at, name
                """
            ),
            {"sid": sid},
        ).mappings().all()
    return [dict(r) for r in rows]


def load_alt_map(sid: str) -> Dict[str, str]:
    with engine.begin() as conn:
        rows = conn.execute(
            text(
                """
                SELECT alternative_id::text AS alternative_id, name
                FROM alternatives
                WHERE scenario_id = :sid
                ORDER BY created_at, name
                """
            ),
            {"sid": sid},
        ).mappings().all()
    return {r["name"]: r["alternative_id"] for r in rows}


def normalize_weight_map(weights_map: Dict[str, float], crits: List[str]) -> Dict[str, float]:
    vec = np.array([float(weights_map.get(c, 0.0)) for c in crits], dtype=float)
    if vec.sum() <= 0:
        vec = np.ones(len(crits), dtype=float) / max(1, len(crits))
    else:
        vec = vec / vec.sum()
    return {crits[i]: float(vec[i]) for i in range(len(crits))}


def weighted_input_matrix(matrix_df: pd.DataFrame, weights_map: Dict[str, float]) -> pd.DataFrame:
    if matrix_df is None or matrix_df.empty:
        return pd.DataFrame()
    aligned = pd.Series({c: float(weights_map.get(c, 0.0)) for c in matrix_df.columns})
    return matrix_df.mul(aligned, axis=1)


def get_vft_tables(run_id: str) -> Tuple[pd.DataFrame, pd.DataFrame]:
    vft_data = vft_svc.get_vft_results(run_id, engine)
    util_list = vft_data.get("utilities", [])
    weighted_list = vft_data.get("weighted", [])
    util_df = pd.DataFrame()
    weighted_df = pd.DataFrame()
    if util_list:
        util_df = pd.DataFrame(util_list).pivot(index="alternative_name", columns="criterion_name", values="utility_value")
    if weighted_list:
        weighted_df = pd.DataFrame(weighted_list).pivot(index="alternative_name", columns="criterion_name", values="weighted_utility")
    return util_df, weighted_df


def load_topsis_weighted_long(run_id: str) -> pd.DataFrame:
    with engine.begin() as conn:
        rows = conn.execute(
            text(
                """
                SELECT a.name AS alternative_name, c.name AS criterion_name, v.value AS weighted_value
                FROM topsis_weighted_values v
                JOIN alternatives a ON a.alternative_id = v.alternative_id
                JOIN criteria c ON c.criterion_id = v.criterion_id
                WHERE v.run_id = :rid
                ORDER BY a.name, c.name
                """
            ),
            {"rid": run_id},
        ).mappings().all()
    return pd.DataFrame([dict(r) for r in rows])


def lighten_color(color, amount: float = 0.0):
    try:
        c = np.array(mcolors.to_rgb(color))
    except Exception:
        c = np.array(mcolors.to_rgb("#4472c4"))
    white = np.array([1.0, 1.0, 1.0])
    mixed = c + (white - c) * max(0.0, min(0.85, amount))
    return tuple(mixed)


BASE_COLORS = [
    "#1f77b4", "#ff7f0e", "#2ca02c", "#d62728", "#9467bd",
    "#8c564b", "#e377c2", "#7f7f7f", "#bcbd22", "#17becf",
]


def get_color_map(labels: List[str]) -> Dict[str, str]:
    return {label: BASE_COLORS[i % len(BASE_COLORS)] for i, label in enumerate(labels)}


def fig_to_buf(fig) -> io.BytesIO:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    buf.seek(0)
    plt.close(fig)
    return buf


def render_option(parent_key: str, label: str, default: bool = False):
    selected = st.checkbox(label, value=default, key=f"{parent_key}_sel")
    note = ""
    if selected:
        note = st.text_area("Note for this element", key=f"{parent_key}_note", height=70)
    return selected, note


def pie_weights_chart(weights_map: Dict[str, float]) -> io.BytesIO | None:
    if not weights_map:
        return None
    labels = list(weights_map.keys())
    values = [float(weights_map[k]) for k in labels]
    if sum(values) <= 0:
        return None
    colors = [BASE_COLORS[i % len(BASE_COLORS)] for i in range(len(labels))]
    fig, ax = plt.subplots(figsize=(8.2, 4.6))
    wedges, _, _ = ax.pie(values, autopct="%1.1f%%", startangle=90, colors=colors, textprops={"fontsize": 9})
    ax.axis("equal")
    ax.set_title("Criterion Weight Distribution", fontsize=11, fontweight="bold")
    ax.legend(wedges, labels, title="Criteria", loc="center left", bbox_to_anchor=(1.0, 0.5), fontsize=8)
    return fig_to_buf(fig)


def ranking_bar_chart(scores_df: pd.DataFrame, title: str, color: str) -> io.BytesIO | None:
    if scores_df is None or scores_df.empty:
        return None
    sdf = scores_df.sort_values("rank")
    fig, ax = plt.subplots(figsize=(7.4, 4.0))
    bars = ax.bar(sdf["alternative_name"], sdf["score"], color=color)
    ax.set_title(title, fontsize=11, fontweight="bold")
    ax.set_ylabel("Score")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.xticks(rotation=25, ha="right", fontsize=9)
    for bar, val in zip(bars, sdf["score"]):
        ax.text(bar.get_x() + bar.get_width() / 2, float(val), f"{float(val):.3f}", ha="center", va="bottom", fontsize=8)
    return fig_to_buf(fig)


def heatmap_chart(df: pd.DataFrame, title: str, cmap: str = "Blues") -> io.BytesIO | None:
    if df is None or df.empty:
        return None
    data = df.astype(float).values
    fig, ax = plt.subplots(figsize=(max(6.2, 0.8 * len(df.columns) + 2), max(3.6, 0.42 * len(df.index) + 2)))
    im = ax.imshow(data, aspect="auto", cmap=cmap)
    ax.set_title(title, fontsize=11, fontweight="bold")
    ax.set_xticks(np.arange(len(df.columns)))
    ax.set_xticklabels(df.columns, rotation=35, ha="right", fontsize=8)
    ax.set_yticks(np.arange(len(df.index)))
    ax.set_yticklabels(df.index, fontsize=8)
    cbar = fig.colorbar(im, ax=ax)
    cbar.ax.set_ylabel("Value", rotation=90)
    for i in range(len(df.index)):
        for j in range(len(df.columns)):
            ax.text(j, i, f"{data[i, j]:.2f}", ha="center", va="center", fontsize=7, color="black")
    return fig_to_buf(fig)


def sandbox_weight_controls(section_key: str, crits: List[str], base_weights_map: Dict[str, float]) -> Dict[str, float]:
    base_norm = normalize_weight_map(base_weights_map, crits)
    st.markdown("**Sandbox weights**")
    auto_norm = st.checkbox("Auto-normalize sandbox weights", value=True, key=f"{section_key}_autonorm")
    cols = st.columns(min(4, max(1, len(crits))))
    raw = {}
    for i, c in enumerate(crits):
        with cols[i % len(cols)]:
            raw[c] = st.slider(c, 0.0, 1.0, float(base_norm.get(c, 0.0)), 0.01, key=f"{section_key}_slider_{c}")
    w = np.array([float(raw[c]) for c in crits], dtype=float)
    if auto_norm and w.sum() > 0:
        w = w / w.sum()
    return {crits[i]: float(w[i]) for i in range(len(crits))}


def topsis_sandbox_from_baseline_run(baseline_run_id: str, directions: List[str], sandbox_weights: Dict[str, float]):
    norm_df = topsis_read.get_matrix(baseline_run_id, "normalized")
    if norm_df is None or norm_df.empty:
        return None
    crits = list(norm_df.columns)
    weights_vec = np.array([float(sandbox_weights.get(c, 0.0)) for c in crits], dtype=float)
    if weights_vec.sum() > 0:
        weights_vec = weights_vec / weights_vec.sum()
    artifacts = compute_topsis(norm_df.values.astype(float), weights_vec, directions)
    score_df = pd.DataFrame({
        "alternative_name": list(norm_df.index),
        "score": artifacts.c_star,
    }).sort_values("score", ascending=False).reset_index(drop=True)
    score_df["rank"] = np.arange(1, len(score_df) + 1)
    weighted_df = pd.DataFrame(artifacts.weighted_matrix, index=norm_df.index, columns=norm_df.columns)
    dist_df = pd.DataFrame({
        "alternative_name": list(norm_df.index),
        "s_pos": artifacts.s_pos,
        "s_neg": artifacts.s_neg,
        "c_star": artifacts.c_star,
    }).sort_values("c_star", ascending=False).reset_index(drop=True)
    return score_df, weighted_df, dist_df


def vft_sandbox_from_baseline_run(baseline_run_id: str, sandbox_weights: Dict[str, float]):
    util_df, _ = get_vft_tables(baseline_run_id)
    if util_df is None or util_df.empty:
        return None
    crits = list(util_df.columns)
    weight_series = pd.Series({c: float(sandbox_weights.get(c, 0.0)) for c in crits})
    weighted_df = util_df.mul(weight_series, axis=1)
    total_scores = weighted_df.sum(axis=1).sort_values(ascending=False)
    score_df = total_scores.reset_index()
    score_df.columns = ["alternative_name", "score"]
    score_df["rank"] = np.arange(1, len(score_df) + 1)
    return score_df, weighted_df


def sandbox_compare_chart(base_scores_df: pd.DataFrame, sandbox_scores_df: pd.DataFrame, title: str) -> io.BytesIO | None:
    if base_scores_df.empty or sandbox_scores_df.empty:
        return None
    merge = base_scores_df[["alternative_name", "score"]].rename(columns={"score": "Baseline"}).merge(
        sandbox_scores_df[["alternative_name", "score"]].rename(columns={"score": "Sandbox"}),
        on="alternative_name",
        how="outer",
    ).fillna(0.0)
    x = np.arange(len(merge))
    fig, ax = plt.subplots(figsize=(7.8, 4.2))
    ax.bar(x - 0.18, merge["Baseline"], width=0.36, label="Baseline", color="#1f77b4")
    ax.bar(x + 0.18, merge["Sandbox"], width=0.36, label="Sandbox", color="#ff7f0e")
    ax.set_xticks(x)
    ax.set_xticklabels(merge["alternative_name"], rotation=25, ha="right", fontsize=9)
    ax.set_ylabel("Score")
    ax.set_title(title, fontsize=11, fontweight="bold")
    ax.legend(fontsize=8)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    return fig_to_buf(fig)


def sandbox_rank_shift_chart(base_scores_df: pd.DataFrame, sandbox_scores_df: pd.DataFrame) -> io.BytesIO | None:
    if base_scores_df.empty or sandbox_scores_df.empty:
        return None
    merge = base_scores_df[["alternative_name", "rank"]].rename(columns={"rank": "Baseline"}).merge(
        sandbox_scores_df[["alternative_name", "rank"]].rename(columns={"rank": "Sandbox"}),
        on="alternative_name",
        how="outer",
    )
    fig, ax = plt.subplots(figsize=(7.6, 4.2))
    color_map = get_color_map(merge["alternative_name"].tolist())
    for _, row in merge.iterrows():
        x = [0, 1]
        y = [float(row["Baseline"]), float(row["Sandbox"])]
        ax.plot(x, y, marker="o", linewidth=2, color=color_map[row["alternative_name"]], label=row["alternative_name"])
    ax.set_xticks([0, 1])
    ax.set_xticklabels(["Baseline", "Sandbox"])
    ax.invert_yaxis()
    ax.set_ylabel("Rank (1 = Best)")
    ax.set_title("Rank Shift: Baseline vs Sandbox", fontsize=11, fontweight="bold")
    ax.legend(fontsize=7, bbox_to_anchor=(1.02, 1), loc="upper left")
    return fig_to_buf(fig)


def stacked_contribution_chart(weighted_df: pd.DataFrame, title: str) -> io.BytesIO | None:
    if weighted_df is None or weighted_df.empty:
        return None
    order = weighted_df.sum(axis=1).sort_values(ascending=False).index.tolist()
    plot_df = weighted_df.reindex(order)
    crits = list(plot_df.columns)
    colors = get_color_map(crits)
    fig, ax = plt.subplots(figsize=(8.0, 4.4))
    bottom = np.zeros(len(plot_df))
    for crit in crits:
        vals = plot_df[crit].astype(float).values
        ax.bar(plot_df.index, vals, bottom=bottom, label=crit, color=colors[crit])
        bottom += vals
    ax.set_title(title, fontsize=11, fontweight="bold")
    ax.set_ylabel("Weighted Contribution")
    plt.xticks(rotation=25, ha="right", fontsize=9)
    ax.legend(fontsize=7, bbox_to_anchor=(1.02, 1), loc="upper left")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    return fig_to_buf(fig)


def topsis_distance_chart(dist_df: pd.DataFrame, title: str) -> io.BytesIO | None:
    if dist_df is None or dist_df.empty:
        return None
    ddf = dist_df.copy()
    x = np.arange(len(ddf))
    fig, ax = plt.subplots(figsize=(8.0, 4.2))
    ax.bar(x - 0.18, ddf["s_pos"], width=0.36, label="S+", color="#9ecae1")
    ax.bar(x + 0.18, ddf["s_neg"], width=0.36, label="S-", color="#3182bd")
    ax.set_xticks(x)
    ax.set_xticklabels(ddf["alternative_name"], rotation=25, ha="right", fontsize=9)
    ax.set_title(title, fontsize=11, fontweight="bold")
    ax.set_ylabel("Distance")
    ax.legend(fontsize=8)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    return fig_to_buf(fig)


def build_compare_payload(sid: str, pref_ids: List[str], method: str, pref_name_map: Dict[str, str]) -> dict:
    long_frames = []
    weighted_payload = []
    dist_payload = []
    valid_labels = []
    for pid in pref_ids:
        run_id = load_latest_run_for_pref(sid, pid, method)
        if not run_id:
            continue
        label = pref_name_map.get(pid, pid)
        valid_labels.append(label)
        sdf = pd.DataFrame(result_repo.get_scores_with_names(run_id))
        if not sdf.empty:
            sdf = sdf.copy()
            sdf["comparison_label"] = label
            long_frames.append(sdf)
        if method == "topsis":
            weighted_long = load_topsis_weighted_long(run_id)
            if not weighted_long.empty:
                weighted_long = weighted_long.rename(columns={
                    "alternative_name": "Alternative",
                    "criterion_name": "Criterion",
                    "weighted_value": "Weighted Value",
                })
                weighted_long["comparison_label"] = label
                weighted_payload.append(weighted_long)
            dist_df = topsis_read.get_distances(run_id)
            if not dist_df.empty:
                dist_df = dist_df.rename(columns={"alternative": "alternative_name"})
                dist_df["comparison_label"] = label
                dist_payload.append(dist_df)
        else:
            _, wt_df = get_vft_tables(run_id)
            if not wt_df.empty:
                long_w = wt_df.stack().reset_index()
                long_w.columns = ["Alternative", "Criterion", "Weighted Value"]
                long_w["comparison_label"] = label
                weighted_payload.append(long_w)
    return {
        "long_df": pd.concat(long_frames, ignore_index=True) if long_frames else pd.DataFrame(),
        "weighted_long_df": pd.concat(weighted_payload, ignore_index=True) if weighted_payload else pd.DataFrame(),
        "dist_df": pd.concat(dist_payload, ignore_index=True) if dist_payload else pd.DataFrame(),
        "selection_order": valid_labels,
        "method": method,
    }


def comparison_score_chart(long_df: pd.DataFrame, selection_order: List[str]) -> io.BytesIO | None:
    if long_df.empty:
        return None
    alt_order = long_df[long_df["comparison_label"] == selection_order[0]].sort_values("rank")["alternative_name"].tolist() if selection_order else sorted(long_df["alternative_name"].unique().tolist())
    pivot = long_df.pivot(index="alternative_name", columns="comparison_label", values="score").reindex(index=alt_order)
    x = np.arange(len(pivot.index))
    fig, ax = plt.subplots(figsize=(8.4, 4.5))
    n = max(1, len(pivot.columns))
    width = 0.8 / n
    for i, col in enumerate(pivot.columns):
        vals = pivot[col].fillna(0.0).values
        ax.bar(x - 0.4 + width / 2 + i * width, vals, width=width, label=col, color=lighten_color("#1f77b4", i * 0.18))
    ax.set_xticks(x)
    ax.set_xticklabels(pivot.index, rotation=25, ha="right", fontsize=9)
    ax.set_title("Alternative Scores Across Selected Preference Sets", fontsize=11, fontweight="bold")
    ax.set_ylabel("Score")
    ax.legend(fontsize=7, bbox_to_anchor=(1.02, 1), loc="upper left")
    return fig_to_buf(fig)


def comparison_rank_shift_chart(long_df: pd.DataFrame, selection_order: List[str]) -> io.BytesIO | None:
    if long_df.empty:
        return None
    alt_names = sorted(long_df["alternative_name"].unique().tolist())
    color_map = get_color_map(alt_names)
    fig, ax = plt.subplots(figsize=(8.2, 4.4))
    for alt in alt_names:
        adf = long_df[long_df["alternative_name"] == alt][["comparison_label", "rank"]].drop_duplicates().set_index("comparison_label").reindex(selection_order).reset_index()
        if adf.empty:
            continue
        ax.plot(adf["comparison_label"], adf["rank"], marker="o", linewidth=2, color=color_map[alt], label=alt)
    ax.invert_yaxis()
    ax.set_ylabel("Rank (1 = Best)")
    ax.set_title("Rank Shift Across Selected Preference Sets", fontsize=11, fontweight="bold")
    ax.legend(fontsize=7, bbox_to_anchor=(1.02, 1), loc="upper left")
    return fig_to_buf(fig)


def comparison_stacked_chart(weighted_long_df: pd.DataFrame, selection_order: List[str]) -> io.BytesIO | None:
    if weighted_long_df.empty:
        return None
    alt_order = weighted_long_df["Alternative"].drop_duplicates().tolist()
    crit_order = weighted_long_df["Criterion"].drop_duplicates().tolist()
    crit_colors = get_color_map(crit_order)
    shade_map = {lab: i * 0.18 for i, lab in enumerate(selection_order)}
    fig, ax = plt.subplots(figsize=(9.0, 4.8))
    group_width = 0.82
    bar_width = group_width / max(1, len(selection_order))
    x = np.arange(len(alt_order))
    for s_idx, comp_label in enumerate(selection_order):
        subset = weighted_long_df[weighted_long_df["comparison_label"] == comp_label]
        xpos = x - group_width / 2 + bar_width / 2 + s_idx * bar_width
        bottom = np.zeros(len(alt_order))
        for crit in crit_order:
            vals = []
            crit_df = subset[subset["Criterion"] == crit]
            for alt in alt_order:
                row = crit_df[crit_df["Alternative"] == alt]
                vals.append(float(row["Weighted Value"].iloc[0]) if not row.empty else 0.0)
            ax.bar(xpos, vals, width=bar_width, bottom=bottom, color=lighten_color(crit_colors[crit], shade_map[comp_label]), label=crit if s_idx == 0 else None)
            bottom += np.array(vals)
    ax.set_xticks(x)
    ax.set_xticklabels(alt_order, rotation=25, ha="right", fontsize=9)
    ax.set_title("Alternative Score Composition Across Selected Sets", fontsize=11, fontweight="bold")
    ax.set_ylabel("Weighted Contribution")
    ax.legend(fontsize=7, bbox_to_anchor=(1.02, 1), loc="upper left", title="Criterion")
    return fig_to_buf(fig)


def comparison_distance_chart(dist_df: pd.DataFrame, selection_order: List[str]) -> io.BytesIO | None:
    if dist_df.empty:
        return None
    alt_order = dist_df[dist_df["comparison_label"] == selection_order[0]]["alternative_name"].tolist() if selection_order else dist_df["alternative_name"].drop_duplicates().tolist()
    fig, ax = plt.subplots(figsize=(9.0, 4.6))
    x = np.arange(len(alt_order))
    n = max(1, len(selection_order))
    width = 0.75 / (2 * n)
    for i, label in enumerate(selection_order):
        sub = dist_df[dist_df["comparison_label"] == label]
        s_pos = [float(sub[sub["alternative_name"] == alt]["s_pos"].iloc[0]) if not sub[sub["alternative_name"] == alt].empty else 0.0 for alt in alt_order]
        s_neg = [float(sub[sub["alternative_name"] == alt]["s_neg"].iloc[0]) if not sub[sub["alternative_name"] == alt].empty else 0.0 for alt in alt_order]
        base_x = x - 0.375 + i * 2 * width
        ax.bar(base_x, s_pos, width=width, color=lighten_color("#1f77b4", i * 0.16), label=f"{label} | S+")
        ax.bar(base_x + width, s_neg, width=width, color=lighten_color("#ff7f0e", i * 0.16), label=f"{label} | S-")
    ax.set_xticks(x)
    ax.set_xticklabels(alt_order, rotation=25, ha="right", fontsize=9)
    ax.set_title("TOPSIS Distance Decomposition Across Selected Sets", fontsize=11, fontweight="bold")
    ax.set_ylabel("Distance")
    ax.legend(fontsize=6.5, bbox_to_anchor=(1.02, 1), loc="upper left")
    return fig_to_buf(fig)


# Selection
st.subheader("Select Data to Report")

decisions = load_decisions()
if not decisions:
    st.warning("No decisions found.")
    st.stop()

dec_ids = [d["decision_id"] for d in decisions]
dec_id_to_title = {d["decision_id"]: d["title"] for d in decisions}
default_dec = st.session_state.get("decision_id") or dec_ids[0]
picked_dec = st.multiselect("Decision", dec_ids, default=[default_dec] if default_dec in dec_ids else [dec_ids[0]], max_selections=1, format_func=lambda x: dec_id_to_title.get(x, x), key="rpt_dec")
decision_id = picked_dec[0] if picked_dec else dec_ids[0]
st.session_state["decision_id"] = decision_id

scenarios = load_scenarios(decision_id)
if not scenarios:
    st.warning("No scenarios.")
    st.stop()
scen_ids = [s["scenario_id"] for s in scenarios]
scen_id_to_name = {s["scenario_id"]: s["name"] for s in scenarios}
scen_id_to_method = {s["scenario_id"]: s.get("method_type", "topsis") for s in scenarios}
default_scen = st.session_state.get("scenario_id")
if default_scen not in scen_ids:
    default_scen = scen_ids[0]
picked_scen = st.multiselect("Scenario", scen_ids, default=[default_scen], max_selections=1, format_func=lambda x: scen_id_to_name.get(x, x), key="rpt_scen")
scenario_id = picked_scen[0] if picked_scen else scen_ids[0]
st.session_state["scenario_id"] = scenario_id
st.session_state["method_choice"] = scen_id_to_method.get(scenario_id, "topsis")
method_choice = st.session_state["method_choice"]

prefs = load_prefs(scenario_id)
if not prefs:
    st.warning("No preference sets.")
    st.stop()
pref_ids = [p["preference_set_id"] for p in prefs]
pref_id_to_name = {p["preference_set_id"]: p["name"] for p in prefs}
default_pref = st.session_state.get("preference_set_id") or pref_ids[0]
if default_pref not in pref_ids:
    default_pref = pref_ids[0]
picked_pref = st.multiselect("Preference set", pref_ids, default=[default_pref], max_selections=1, format_func=lambda x: pref_id_to_name.get(x, x), key="rpt_pref")
pref_id = picked_pref[0] if picked_pref else pref_ids[0]
st.session_state["preference_set_id"] = pref_id

runs = load_runs(scenario_id, pref_id)
if not runs:
    st.warning("No runs found.")
    st.stop()

run_ids = [r["run_id"] for r in runs]
run_id_to_row = {r["run_id"]: r for r in runs}


def run_fmt(rid):
    r = run_id_to_row.get(rid, {})
    lbl = (r.get("run_label") or "").strip()
    by = f" · {r['executed_by']}" if r.get("executed_by") else ""
    ts = str(r.get("executed_at", ""))[:16]
    tag = r.get("method", "").upper()
    base = f"[{tag}] {ts}{by}"
    return f"{lbl} | {base}" if lbl else base


default_run = st.session_state.get("last_run_id")
if default_run not in run_ids:
    default_run = run_ids[0]

picked_run = st.multiselect("Run to report", run_ids, default=[default_run], max_selections=1, format_func=run_fmt, key=f"rpt_runs_{scenario_id}_{pref_id}")
if not picked_run:
    st.info("Select a run.")
    st.stop()

run_id = picked_run[0]
st.session_state["last_run_id"] = run_id
run_meta = run_id_to_row.get(run_id, {})
run_method = run_meta.get("method", method_choice)

st.divider()

# Header
st.subheader("Report Header")
col_title, col_author = st.columns(2)
with col_title:
    report_title = st.text_input("Report title", value=st.session_state.get("report_title", "MCDA Analysis Report"), key="rpt_title_input")
    st.session_state["report_title"] = report_title
with col_author:
    author = st.text_input("Author", value=user_name, key="rpt_author")

exec_summary = st.text_area(
    "Executive Summary",
    value=st.session_state.get("rpt_exec_summary", ""),
    height=100,
    placeholder="Key findings, context, or recommendations for this analysis.",
    key="rpt_exec_text",
)
st.session_state["rpt_exec_summary"] = exec_summary

st.divider()

# Load data
matrix_df = meas_repo.load_matrix_ui(scenario_id)
criteria_meta = load_criteria_meta(scenario_id)
crit_names = [c["name"] for c in criteria_meta]
weights_map = normalize_weight_map(pref_repo.load_weights_by_criterion_name(pref_id), crit_names)
scores_df = pd.DataFrame(result_repo.get_scores_with_names(run_id))
if not scores_df.empty:
    scores_df = scores_df.sort_values("rank").reset_index(drop=True)
weighted_input_df = weighted_input_matrix(matrix_df, weights_map)
util_df, weighted_util_df = get_vft_tables(run_id) if run_method == "vft" else (pd.DataFrame(), pd.DataFrame())

# Per-section content
st.subheader("Section Content, Elements, and Notes")
sections: dict = {}

with st.expander("Section 1: Input Matrix and Weights", expanded=True):
    sections["sec1_enable"] = st.checkbox("Include Section 1", value=True, key="sec1_enable")
    if sections["sec1_enable"]:
        sections["sec1_table"] = render_option("sec1_table", "Select Table", True)
        sections["sec1_pie"] = render_option("sec1_pie", "Select Pie Chart (with legend)")
        sections["sec1_weighted"] = render_option("sec1_weighted", "Select Weighted Final Table")
        if matrix_df is not None and not matrix_df.empty:
            st.dataframe(matrix_df, use_container_width=True)
            if not weighted_input_df.empty:
                st.dataframe(weighted_input_df.style.format("{:.4f}"), use_container_width=True)

with st.expander("Section 2: Ranking Chart and Table", expanded=True):
    sections["sec2_enable"] = st.checkbox("Include Section 2", value=True, key="sec2_enable")
    if sections["sec2_enable"]:
        sections["sec2_table"] = render_option("sec2_table", "Select Final Scoring Table", True)
        sections["sec2_bar"] = render_option("sec2_bar", "Select Bar Graph of the Ranking", True)
        if not scores_df.empty:
            st.dataframe(scores_df, use_container_width=True)

if run_method == "vft":
    with st.expander("Section 3: VFT Utility Matrix & Weighted Matrix", expanded=False):
        sections["sec3_enable"] = st.checkbox("Include Section 3", value=True, key="sec3_enable")
        if sections["sec3_enable"]:
            sections["sec3_util_heat"] = render_option("sec3_util_heat", "Select heatmap for Utility Matrix (with legend)")
            sections["sec3_wutil_heat"] = render_option("sec3_wutil_heat", "Select heatmap for Weighted Utility Matrix (with legend)")
            sections["sec3_util_table"] = render_option("sec3_util_table", "Select the table for Utility Matrix", True)
            sections["sec3_wutil_table"] = render_option("sec3_wutil_table", "Select table for Weighted Utility Matrix", True)
            if not util_df.empty:
                st.dataframe(util_df.style.format("{:.4f}"), use_container_width=True)
            if not weighted_util_df.empty:
                st.dataframe(weighted_util_df.style.format("{:.4f}"), use_container_width=True)

with st.expander("Section 4: Sensitivity Analysis", expanded=False):
    sections["sec4_enable"] = st.checkbox("Include Section 4", value=False, key="sec4_enable")
    if sections["sec4_enable"]:
        sb_pref_pick = st.multiselect(
            "Baseline preference set",
            options=pref_ids,
            default=[pref_id],
            max_selections=1,
            format_func=lambda x: pref_id_to_name.get(x, x),
            key="rpt_sb_pref_pick",
        )
        sb_pref_id = sb_pref_pick[0] if sb_pref_pick else pref_id
        base_weights_for_sb = normalize_weight_map(pref_repo.load_weights_by_criterion_name(sb_pref_id), crit_names)
        sandbox_weights = sandbox_weight_controls("sec4_sb", crit_names, base_weights_for_sb)
        base_weights_df = pd.DataFrame([base_weights_for_sb], index=["Baseline"])
        sandbox_weights_df = pd.DataFrame([sandbox_weights], index=["Sandbox"])
        sections["sec4_base_table"] = render_option("sec4_base_table", "Include baseline weight table", True)
        sections["sec4_sb_table"] = render_option("sec4_sb_table", "Include sandbox weight table", True)
        sections["sec4_score_cmp"] = render_option("sec4_score_cmp", "Include baseline vs sandbox score chart", True)
        sections["sec4_rank_shift"] = render_option("sec4_rank_shift", "Include baseline vs sandbox rank shift graph", True)
        sections["sec4_contrib"] = render_option("sec4_contrib", "Include sandbox contribution chart", run_method == "vft")
        if run_method == "topsis":
            sections["sec4_dist"] = render_option("sec4_dist", "Include TOPSIS distance decomposition chart")
        st.dataframe(pd.concat([base_weights_df, sandbox_weights_df]).style.format("{:.4f}"), use_container_width=True)

        baseline_run_id = load_latest_run_for_pref(scenario_id, sb_pref_id, run_method)
        sensitivity_payload = None
        if baseline_run_id:
            if run_method == "topsis":
                directions = [(c.get("direction") or "benefit").strip().lower() for c in criteria_meta]
                sensitivity_payload = topsis_sandbox_from_baseline_run(baseline_run_id, directions, sandbox_weights)
            else:
                sensitivity_payload = vft_sandbox_from_baseline_run(baseline_run_id, sandbox_weights)

        if sensitivity_payload:
            if run_method == "topsis":
                sb_scores_df, sb_weighted_df, sb_dist_df = sensitivity_payload
            else:
                sb_scores_df, sb_weighted_df = sensitivity_payload
                sb_dist_df = pd.DataFrame()
            st.dataframe(sb_scores_df, use_container_width=True)
            if not sb_weighted_df.empty:
                st.dataframe(sb_weighted_df.style.format("{:.4f}"), use_container_width=True)

with st.expander("Section 5: Compared Preference Sets", expanded=False):
    sections["sec5_enable"] = st.checkbox("Include Section 5", value=False, key="sec5_enable")
    if sections["sec5_enable"]:
        default_cmp = pref_ids[: min(2, len(pref_ids))]
        cmp_pref_ids = st.multiselect(
            "Preference sets to compare (2-5)",
            options=pref_ids,
            default=default_cmp,
            max_selections=5,
            format_func=lambda x: pref_id_to_name.get(x, x),
            key="rpt_cmp_pref_ids",
        )
        sections["sec5_score"] = render_option("sec5_score", "Include score comparison graph", True)
        sections["sec5_rank"] = render_option("sec5_rank", "Include rank shift graph", True)
        sections["sec5_stack"] = render_option("sec5_stack", "Include stacked contribution comparison graph", True)
        if run_method == "topsis":
            sections["sec5_dist"] = render_option("sec5_dist", "Include TOPSIS distance comparison graph")
        if len(cmp_pref_ids) >= 2:
            cmp_payload = build_compare_payload(scenario_id, cmp_pref_ids, run_method, pref_id_to_name)
            sections["cmp_payload"] = cmp_payload
            if not cmp_payload["long_df"].empty:
                st.dataframe(cmp_payload["long_df"], use_container_width=True)
        else:
            st.info("Select at least 2 preference sets to enable this section.")

st.divider()
st.subheader("Export Report")
col_e1, col_e2 = st.columns([1, 3])
with col_e1:
    do_export = st.button("Generate Report (DOCX)", type="primary", key="btn_gen_report")
with col_e2:
    st.caption("Generates a Word document with the selected sections, charts, tables, legends, and notes.")

if do_export:
    try:
        from docx import Document as DocxDoc
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = DocxDoc()
        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Arial"
        normal_style.font.size = Pt(10)

        def add_note(doc_obj, note_text, label="Note"):
            if note_text and str(note_text).strip():
                p = doc_obj.add_paragraph()
                r1 = p.add_run(f"{label}: ")
                r1.bold = True
                r1.font.color.rgb = RGBColor(43, 108, 176)
                r2 = p.add_run(str(note_text).strip())
                r2.italic = True
                r2.font.color.rgb = RGBColor(74, 85, 104)

        def add_caption(doc_obj, text_val: str):
            p = doc_obj.add_paragraph(text_val)
            if p.runs:
                p.runs[0].font.size = Pt(8)
                p.runs[0].italic = True

        def df_to_table(doc_obj, df: pd.DataFrame, index: bool = True):
            display_df = df.reset_index() if index else df.reset_index(drop=True)
            table = doc_obj.add_table(rows=1 + len(display_df), cols=len(display_df.columns))
            table.style = "Table Grid"
            for j, col in enumerate(display_df.columns):
                cell = table.rows[0].cells[j]
                cell.text = str(col)
                if cell.paragraphs and cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].bold = True
            for i, row in display_df.iterrows():
                for j, val in enumerate(row):
                    try:
                        num = float(val)
                        table.rows[i + 1].cells[j].text = f"{num:.4f}"
                    except Exception:
                        table.rows[i + 1].cells[j].text = str(val)

        def add_picture(doc_obj, buf, width=6.2):
            if buf is not None:
                doc_obj.add_picture(buf, width=Inches(width))

        title_p = doc.add_heading(report_title, level=0)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Author: {author or '—'}")
        doc.add_paragraph(f"Decision: {dec_id_to_title.get(decision_id, decision_id)}")
        doc.add_paragraph(f"Scenario: {scen_id_to_name.get(scenario_id, scenario_id)}")
        doc.add_paragraph(f"Preference Set: {pref_id_to_name.get(pref_id, pref_id)}")
        doc.add_paragraph(f"Method: {run_method.upper()}")
        doc.add_paragraph(f"Run: {run_fmt(run_id)}")
        doc.add_paragraph("")

        if exec_summary.strip():
            doc.add_heading("Executive Summary", level=1)
            for line in exec_summary.strip().split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())
            doc.add_paragraph("")

        if sections.get("sec1_enable"):
            doc.add_heading("1. Input Matrix and Weights", level=1)
            selected, note = sections.get("sec1_table", (False, ""))
            if selected and matrix_df is not None and not matrix_df.empty:
                doc.add_heading("1.1 Input Matrix Table", level=2)
                df_to_table(doc, matrix_df, index=True)
                add_caption(doc, "Table 1.1: Input performance matrix.")
                add_note(doc, note)
            selected, note = sections.get("sec1_pie", (False, ""))
            if selected:
                doc.add_heading("1.2 Criterion Weight Distribution", level=2)
                add_picture(doc, pie_weights_chart(weights_map), width=6.4)
                add_caption(doc, "Fig 1.2: Pie chart of normalized criterion weights with legend.")
                add_note(doc, note)
            selected, note = sections.get("sec1_weighted", (False, ""))
            if selected and not weighted_input_df.empty:
                doc.add_heading("1.3 Weighted Final Table", level=2)
                df_to_table(doc, weighted_input_df, index=True)
                add_caption(doc, "Table 1.3: Weighted input matrix using the selected preference set weights.")
                add_note(doc, note)

        if sections.get("sec2_enable"):
            doc.add_heading("2. Ranking Chart and Table", level=1)
            selected, note = sections.get("sec2_table", (False, ""))
            if selected and not scores_df.empty:
                doc.add_heading("2.1 Final Scoring Table", level=2)
                df_to_table(doc, scores_df, index=False)
                add_caption(doc, "Table 2.1: Final scores and ranks for the selected run.")
                add_note(doc, note)
            selected, note = sections.get("sec2_bar", (False, ""))
            if selected and not scores_df.empty:
                doc.add_heading("2.2 Ranking Bar Graph", level=2)
                add_picture(doc, ranking_bar_chart(scores_df, f"{'TOPSIS' if run_method == 'topsis' else 'VFT'} Ranking by Alternative", "#1f77b4" if run_method == "topsis" else "#2ca25f"), width=6.2)
                add_caption(doc, "Fig 2.2: Bar graph of final alternative scores.")
                add_note(doc, note)

        if run_method == "vft" and sections.get("sec3_enable"):
            doc.add_heading("3. VFT Utility Matrix and Weighted Matrix", level=1)
            selected, note = sections.get("sec3_util_heat", (False, ""))
            if selected and not util_df.empty:
                doc.add_heading("3.1 Utility Matrix Heatmap", level=2)
                add_picture(doc, heatmap_chart(util_df, "Utility Matrix Heatmap", cmap="Blues"), width=6.5)
                add_caption(doc, "Fig 3.1: Heatmap of VFT utility values with color legend.")
                add_note(doc, note)
            selected, note = sections.get("sec3_wutil_heat", (False, ""))
            if selected and not weighted_util_df.empty:
                doc.add_heading("3.2 Weighted Utility Heatmap", level=2)
                add_picture(doc, heatmap_chart(weighted_util_df, "Weighted Utility Heatmap", cmap="Greens"), width=6.5)
                add_caption(doc, "Fig 3.2: Heatmap of weighted utility values with color legend.")
                add_note(doc, note)
            selected, note = sections.get("sec3_util_table", (False, ""))
            if selected and not util_df.empty:
                doc.add_heading("3.3 Utility Matrix Table", level=2)
                df_to_table(doc, util_df, index=True)
                add_caption(doc, "Table 3.3: Utility matrix table.")
                add_note(doc, note)
            selected, note = sections.get("sec3_wutil_table", (False, ""))
            if selected and not weighted_util_df.empty:
                doc.add_heading("3.4 Weighted Utility Matrix Table", level=2)
                df_to_table(doc, weighted_util_df, index=True)
                add_caption(doc, "Table 3.4: Weighted utility matrix table.")
                add_note(doc, note)

        if sections.get("sec4_enable"):
            doc.add_heading("4. Sensitivity Analysis", level=1)
            sb_pref_pick = st.session_state.get("rpt_sb_pref_pick", [pref_id])
            sb_pref_id = sb_pref_pick[0] if isinstance(sb_pref_pick, list) and sb_pref_pick else pref_id
            base_weights_for_sb = normalize_weight_map(pref_repo.load_weights_by_criterion_name(sb_pref_id), crit_names)
            sandbox_weights = {c: float(st.session_state.get(f"sec4_sb_slider_{c}", base_weights_for_sb.get(c, 0.0))) for c in crit_names}
            if st.session_state.get("sec4_sb_autonorm", True):
                sandbox_weights = normalize_weight_map(sandbox_weights, crit_names)
            baseline_run_id = load_latest_run_for_pref(scenario_id, sb_pref_id, run_method)
            if baseline_run_id:
                base_scores_df = pd.DataFrame(result_repo.get_scores_with_names(baseline_run_id)).sort_values("rank")
                if run_method == "topsis":
                    directions = [(c.get("direction") or "benefit").strip().lower() for c in criteria_meta]
                    sens = topsis_sandbox_from_baseline_run(baseline_run_id, directions, sandbox_weights)
                    sb_scores_df, sb_weighted_df, sb_dist_df = sens if sens else (pd.DataFrame(), pd.DataFrame(), pd.DataFrame())
                else:
                    sens = vft_sandbox_from_baseline_run(baseline_run_id, sandbox_weights)
                    sb_scores_df, sb_weighted_df = sens if sens else (pd.DataFrame(), pd.DataFrame())
                    sb_dist_df = pd.DataFrame()

                selected, note = sections.get("sec4_base_table", (False, ""))
                if selected:
                    doc.add_heading("4.1 Baseline Weight Table", level=2)
                    df_to_table(doc, pd.DataFrame([base_weights_for_sb], index=["Baseline"]), index=True)
                    add_caption(doc, "Table 4.1: Baseline normalized weight vector.")
                    add_note(doc, note)
                selected, note = sections.get("sec4_sb_table", (False, ""))
                if selected:
                    doc.add_heading("4.2 Sandbox Weight Table", level=2)
                    df_to_table(doc, pd.DataFrame([sandbox_weights], index=["Sandbox"]), index=True)
                    add_caption(doc, "Table 4.2: Sandbox normalized weight vector used for sensitivity analysis.")
                    add_note(doc, note)
                selected, note = sections.get("sec4_score_cmp", (False, ""))
                if selected and not base_scores_df.empty and not sb_scores_df.empty:
                    doc.add_heading("4.3 Baseline vs Sandbox Score Chart", level=2)
                    add_picture(doc, sandbox_compare_chart(base_scores_df, sb_scores_df, "Baseline vs Sandbox Score Comparison"), width=6.4)
                    add_caption(doc, "Fig 4.3: Score comparison between baseline and sandbox weights.")
                    add_note(doc, note)
                selected, note = sections.get("sec4_rank_shift", (False, ""))
                if selected and not base_scores_df.empty and not sb_scores_df.empty:
                    doc.add_heading("4.4 Baseline vs Sandbox Rank Shift", level=2)
                    add_picture(doc, sandbox_rank_shift_chart(base_scores_df, sb_scores_df), width=6.2)
                    add_caption(doc, "Fig 4.4: Rank shift between baseline and sandbox weights.")
                    add_note(doc, note)
                selected, note = sections.get("sec4_contrib", (False, ""))
                if selected and not sb_weighted_df.empty:
                    doc.add_heading("4.5 Sandbox Contribution Chart", level=2)
                    add_picture(doc, stacked_contribution_chart(sb_weighted_df, "Sandbox Contribution by Criterion"), width=6.6)
                    add_caption(doc, "Fig 4.5: Stacked contribution chart under sandbox weights.")
                    add_note(doc, note)
                selected, note = sections.get("sec4_dist", (False, ""))
                if selected and not sb_dist_df.empty:
                    doc.add_heading("4.6 Sandbox Distance Decomposition", level=2)
                    add_picture(doc, topsis_distance_chart(sb_dist_df, "Sandbox TOPSIS Distances"), width=6.6)
                    add_caption(doc, "Fig 4.6: TOPSIS distance decomposition under sandbox weights.")
                    add_note(doc, note)

        if sections.get("sec5_enable"):
            cmp_pref_ids = st.session_state.get("rpt_cmp_pref_ids", [])
            cmp_payload = sections.get("cmp_payload") or (build_compare_payload(scenario_id, cmp_pref_ids, run_method, pref_id_to_name) if len(cmp_pref_ids) >= 2 else None)
            if cmp_payload and not cmp_payload["long_df"].empty:
                doc.add_heading("5. Compared Preference Sets", level=1)
                selected, note = sections.get("sec5_score", (False, ""))
                if selected:
                    doc.add_heading("5.1 Score Comparison Graph", level=2)
                    add_picture(doc, comparison_score_chart(cmp_payload["long_df"], cmp_payload["selection_order"]), width=6.6)
                    add_caption(doc, "Fig 5.1: Alternative scores across the selected preference sets.")
                    add_note(doc, note)
                selected, note = sections.get("sec5_rank", (False, ""))
                if selected:
                    doc.add_heading("5.2 Rank Shift Graph", level=2)
                    add_picture(doc, comparison_rank_shift_chart(cmp_payload["long_df"], cmp_payload["selection_order"]), width=6.5)
                    add_caption(doc, "Fig 5.2: Rank movement across the selected preference sets.")
                    add_note(doc, note)
                selected, note = sections.get("sec5_stack", (False, ""))
                if selected and not cmp_payload["weighted_long_df"].empty:
                    doc.add_heading("5.3 Stacked Contribution Comparison", level=2)
                    add_picture(doc, comparison_stacked_chart(cmp_payload["weighted_long_df"], cmp_payload["selection_order"]), width=6.8)
                    add_caption(doc, "Fig 5.3: Grouped stacked contribution chart for the selected preference sets.")
                    add_note(doc, note)
                selected, note = sections.get("sec5_dist", (False, ""))
                if selected and not cmp_payload["dist_df"].empty:
                    doc.add_heading("5.4 TOPSIS Distance Comparison", level=2)
                    add_picture(doc, comparison_distance_chart(cmp_payload["dist_df"], cmp_payload["selection_order"]), width=6.8)
                    add_caption(doc, "Fig 5.4: TOPSIS distance decomposition across the selected preference sets.")
                    add_note(doc, note)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        st.download_button(
            "Download Report (.docx)",
            data=buf.getvalue(),
            file_name=f"mcda_report_{scenario_id[:8]}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="rpt_download_btn",
        )
        st.success("Report generated. Click the download button above.")

    except ImportError:
        st.error("python-docx is required. Install it with: pip install python-docx")
    except Exception as e:
        st.error(f"Report generation failed: {e}")
        import traceback
        st.code(traceback.format_exc())
