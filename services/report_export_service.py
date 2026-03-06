# services/report_export_service.py
from __future__ import annotations

from io import BytesIO
from typing import Any, Dict, List

import pandas as pd

from sqlalchemy import text

from docx import Document
from docx.shared import Inches

import matplotlib.pyplot as plt


class ReportExportService:
    """
    Generates a DOCX report for selected runs.

    This service:
    - Reads run outputs (ranking) from DB via existing tables
    - Renders lightweight matplotlib charts (no kaleido required)
    - Inserts tables and charts into a Word document
    """

    def __init__(self, engine):
        self.engine = engine

    # ----------------------------
    # Public
    # ----------------------------
    def build_docx_report(self, payload: Dict[str, Any]) -> bytes:
        doc = Document()

        title = (payload.get("report_title") or "MCDA Run Report").strip()
        author = (payload.get("author") or "").strip()

        doc.add_heading(title, level=0)
        if author:
            doc.add_paragraph(f"Author: {author}")

        doc.add_paragraph(
            f"Decision: {payload.get('decision_title','')} ({payload.get('decision_id','')})"
        )
        doc.add_paragraph(
            f"Scenario: {payload.get('scenario_name','')} ({payload.get('scenario_id','')})"
        )
        doc.add_paragraph(
            f"Preference set: {payload.get('preference_set_name','')} ({payload.get('preference_set_id','')})"
        )

        notes = (payload.get("notes") or "").strip()
        if notes:
            doc.add_heading("Notes", level=1)
            for para in notes.split("\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())

        doc.add_heading("Inputs used for results", level=1)
        inputs_df: pd.DataFrame = payload["inputs_table"]
        self._add_df_table(doc, inputs_df, float_fmt="{:.4f}")

        # Per-run sections
        runs: List[Dict[str, Any]] = payload.get("runs", [])
        if not runs:
            doc.add_paragraph("No runs selected.")
            return self._to_bytes(doc)

        doc.add_heading("Run summaries", level=1)

        all_scores_for_compare = []
        for r in runs:
            run_id = r["run_id"]
            doc.add_heading(r.get("run_label_display", run_id), level=2)

            meta_lines = [
                f"run_id: {run_id}",
                f"method: {r.get('method','')}",
                f"executed_at: {r.get('executed_at','')}",
                f"executed_by: {r.get('executed_by','')}",
                f"engine_version: {r.get('engine_version','')}",
            ]
            rl = (r.get("run_label") or "").strip()
            if rl:
                meta_lines.append(f"run_label: {rl}")
            doc.add_paragraph("\n".join(meta_lines))

            scores_df = self._load_scores(run_id)
            if scores_df.empty:
                doc.add_paragraph("No scores found for this run.")
                continue

            # Table
            doc.add_paragraph("Ranking")
            self._add_df_table(doc, scores_df, float_fmt="{:.6f}")

            # Chart image
            img = self._chart_scores(scores_df, title="TOPSIS score (C*) by alternative")
            doc.add_picture(img, width=Inches(6.5))

            # for comparison chart
            tmp = scores_df.copy()
            tmp["run_id"] = run_id
            tmp["run_label_display"] = r.get("run_label_display", run_id[:8])
            all_scores_for_compare.append(tmp)

        # Comparison section
        if len(all_scores_for_compare) >= 2:
            doc.add_heading("Comparison across runs", level=1)
            comp_df = pd.concat(all_scores_for_compare, axis=0, ignore_index=True)

            # Comparison chart: one line per run
            img2 = self._chart_compare(comp_df)
            doc.add_picture(img2, width=Inches(6.5))

        return self._to_bytes(doc)

    # ----------------------------
    # DB reads
    # ----------------------------
    def _load_scores(self, run_id: str) -> pd.DataFrame:
        """
        Uses result_scores joined to alternatives to return:
        alternative_name, score, rank
        """
        with self.engine.begin() as conn:
            rows = conn.execute(
                text(
                    """
                    SELECT a.name AS alternative_name,
                           rs.score AS score,
                           rs.rank  AS rank
                    FROM result_scores rs
                    JOIN alternatives a ON a.alternative_id = rs.alternative_id
                    WHERE rs.run_id = :rid
                    ORDER BY rs.rank ASC
                    """
                ),
                {"rid": run_id},
            ).mappings().all()

        df = pd.DataFrame([dict(r) for r in rows])
        if not df.empty:
            # ensure types
            df["score"] = pd.to_numeric(df["score"], errors="coerce")
            df["rank"] = pd.to_numeric(df["rank"], errors="coerce").astype("Int64")
        return df

    # ----------------------------
    # DOC helpers
    # ----------------------------
    def _add_df_table(self, doc: Document, df: pd.DataFrame, float_fmt: str = "{:.4f}") -> None:
        if df is None or df.empty:
            doc.add_paragraph("No data.")
            return

        # Include index as first column
        df2 = df.copy()
        df2.insert(0, "Row", df2.index.astype(str))

        table = doc.add_table(rows=1, cols=len(df2.columns))
        hdr = table.rows[0].cells
        for j, col in enumerate(df2.columns):
            hdr[j].text = str(col)

        for _, row in df2.iterrows():
            cells = table.add_row().cells
            for j, col in enumerate(df2.columns):
                val = row[col]
                cells[j].text = self._format_cell(val, float_fmt)

    def _format_cell(self, val: Any, float_fmt: str) -> str:
        if pd.isna(val):
            return ""
        try:
            if isinstance(val, (float, int)) and not isinstance(val, bool):
                return float_fmt.format(float(val))
        except Exception:
            pass
        return str(val)

    def _to_bytes(self, doc: Document) -> bytes:
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio.read()

    # ----------------------------
    # Chart helpers (matplotlib)
    # ----------------------------
    def _chart_scores(self, scores_df: pd.DataFrame, title: str) -> BytesIO:
        df = scores_df.sort_values("rank", ascending=True).copy()
        x = df["alternative_name"].astype(str).tolist()
        y = df["score"].astype(float).tolist()

        fig = plt.figure(figsize=(9, 3.8))
        ax = fig.add_subplot(111)
        ax.bar(x, y)
        ax.set_title(title)
        ax.set_xlabel("Alternative")
        ax.set_ylabel("Score (C*)")
        ax.tick_params(axis="x", rotation=20)

        fig.tight_layout()
        out = BytesIO()
        fig.savefig(out, format="png", dpi=200)
        plt.close(fig)
        out.seek(0)
        return out

    def _chart_compare(self, comp_df: pd.DataFrame) -> BytesIO:
        """
        Line chart: x=alternative, y=score, one line per run.
        """
        fig = plt.figure(figsize=(9, 4.2))
        ax = fig.add_subplot(111)

        # stable alternative order
        alts = sorted(comp_df["alternative_name"].astype(str).unique().tolist())
        x_idx = list(range(len(alts)))

        for run_label, sub in comp_df.groupby("run_label_display"):
            sub2 = sub.set_index(sub["alternative_name"].astype(str))["score"]
            y = [float(sub2.get(a, float("nan"))) for a in alts]
            ax.plot(x_idx, y, marker="o", label=str(run_label))

        ax.set_title("Score comparison across runs")
        ax.set_xlabel("Alternative")
        ax.set_ylabel("Score (C*)")
        ax.set_xticks(x_idx)
        ax.set_xticklabels(alts, rotation=20)
        ax.legend(loc="best", fontsize=8)

        fig.tight_layout()
        out = BytesIO()
        fig.savefig(out, format="png", dpi=200)
        plt.close(fig)
        out.seek(0)
        return out
